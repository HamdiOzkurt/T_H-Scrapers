import os
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib.dates as mdates
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO
from datetime import datetime
import numpy as np
import re
import traceback
from PyQt6.QtCore import QObject, pyqtSignal, pyqtSlot

# Proje konfigürasyonunu import et
from config import OUTPUT_DIR, REPORT_FILE_TEMPLATE

# --- YENİ EKLENEN IMPORT ---
try:
    import ollama
except ImportError:
    print("UYARI: 'ollama' kütüphanesi bulunamadı. Lütfen 'pip install ollama' ile yükleyin. LLM özellikleri çalışmayacak.")
    ollama = None

# Matplotlib thread güvenliği
import matplotlib
matplotlib.use('Agg')

class ReportBuilderWorker(QObject):
    finished = pyqtSignal(str)
    error = pyqtSignal(str)
    status_message = pyqtSignal(str)
    charts_generated = pyqtSignal(dict)
    progress_updated = pyqtSignal(int, int)  # current_step, total_steps

    def __init__(self, csv_file_path, product_name, category_definitions, product_id,
                 llm_model_name="gemma3:4b", llm_temperature=0.4, charts_only=False):
        super().__init__()
        self.csv_file_path = csv_file_path
        self.product_name = product_name
        self.category_definitions = category_definitions
        self.product_id = product_id
        self._is_running = False
        self.llm_model_name = llm_model_name
        self.llm_temperature = llm_temperature
        self.df_comments = None
        self.charts_only = charts_only
        self._stop_requested = False
        
        self.STYLE_CONFIG = {
            'figure_facecolor': '#1a1a2e',
            'axes_facecolor': '#2c3e50',
            'text_color': '#E4E6EA',
            'grid_color': '#555555',
            'positive_color': '#27ae60',
            'negative_color': '#e74c3c',
            'neutral_color': '#f39c12',
        }

    def stop(self):
        try:
            self._stop_requested = True
            self._is_running = False
            self.status_message.emit("Rapor oluşturma durduruluyor...")
        except Exception:
            pass

    def _apply_dark_theme(self, fig, ax):
        fig.patch.set_facecolor(self.STYLE_CONFIG['figure_facecolor'])
        ax.set_facecolor(self.STYLE_CONFIG['axes_facecolor'])
        ax.tick_params(axis='x', colors=self.STYLE_CONFIG['text_color'])
        ax.tick_params(axis='y', colors=self.STYLE_CONFIG['text_color'])
        ax.spines['bottom'].set_color(self.STYLE_CONFIG['grid_color'])
        ax.spines['left'].set_color(self.STYLE_CONFIG['grid_color'])
        ax.spines['top'].set_visible(False)
        ax.spines['right'].set_visible(False)
        ax.title.set_color(self.STYLE_CONFIG['text_color'])
        ax.xaxis.label.set_color(self.STYLE_CONFIG['text_color'])
        ax.yaxis.label.set_color(self.STYLE_CONFIG['text_color'])
        
        if ax.get_legend():
            legend = ax.get_legend()
            legend.get_frame().set_facecolor(self.STYLE_CONFIG['axes_facecolor'])
            legend.get_frame().set_edgecolor(self.STYLE_CONFIG['grid_color'])
            for text in legend.get_texts():
                text.set_color(self.STYLE_CONFIG['text_color'])

    def _load_and_preprocess_data(self):
        try:
            if not os.path.exists(self.csv_file_path):
                self.error.emit(f"Girdi dosyası bulunamadı: {self.csv_file_path}")
                return False
            self.df_comments = pd.read_csv(self.csv_file_path, encoding='utf-8', low_memory=False)
            if 'hesaplanan_tarih' in self.df_comments.columns:
                self.df_comments['hesaplanan_tarih'] = pd.to_datetime(self.df_comments['hesaplanan_tarih'], errors='coerce')
                self.df_comments.dropna(subset=['hesaplanan_tarih'], inplace=True)
            return True
        except Exception as e:
            self.error.emit(f"Veri yükleme hatası: {e}")
            return False

    def _create_sentiment_pie_chart(self):
        sentiment_counts = self.df_comments['duygu_tahmini'].value_counts()
        labels = sentiment_counts.index
        sizes = sentiment_counts.values
        colors_dict = {
            'positive': self.STYLE_CONFIG['positive_color'],
            'negative': self.STYLE_CONFIG['negative_color'],
            'neutral': self.STYLE_CONFIG['neutral_color']
        }
        colors = [colors_dict.get(str(label).lower(), '#bdc3c7') for label in labels]
        if sum(sizes) == 0: return None
        fig, ax = plt.subplots(figsize=(6, 5)) 
        fig.patch.set_facecolor(self.STYLE_CONFIG['figure_facecolor'])
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90,
               colors=colors,
               wedgeprops=dict(width=0.4, edgecolor=self.STYLE_CONFIG['figure_facecolor']),
               pctdistance=0.8, 
               textprops={'color': self.STYLE_CONFIG['text_color'], 'fontsize': 12, 'weight': 'bold'})
        for pct in autotexts:
            pct.set_color("white")
        ax.set_title('Yorumların Genel Duygu Dağılımı', fontsize=14, pad=20, color=self.STYLE_CONFIG['text_color'])
        ax.axis('equal')
        plt.tight_layout(pad=1.5)
        buffer = BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight', dpi=100, facecolor=fig.get_facecolor())
        buffer.seek(0)
        plt.close(fig)
        return buffer

    def _create_category_pie_chart(self):
        available_category_cols = [col for col in self.df_comments.columns
                                  if col not in ['text', 'duygu_tahmini', 'duygu_skoru', 'hesaplanan_tarih', 'date']
                                  and self.df_comments[col].dtype in ['int64', 'float64']]
        category_counts = {}
        for cat in self.category_definitions:
            if cat['csv_col'] in self.df_comments.columns:
                count = self.df_comments[cat['csv_col']].sum()
                if count > 0:
                    category_counts[cat['display_name']] = count
        if not category_counts or sum(category_counts.values()) == 0:
            for col in available_category_cols:
                if col not in [c['csv_col'] for c in self.category_definitions]:
                    count = self.df_comments[col].sum()
                    if count > 0:
                        display_name = col.replace('_', ' ').title()
                        category_counts[display_name] = count
            if not category_counts: return None
        labels = list(category_counts.keys())
        sizes = list(category_counts.values())
        colors = plt.cm.get_cmap('viridis')(np.linspace(0, 1, len(labels)))
        fig, ax = plt.subplots(figsize=(7, 5))
        fig.patch.set_facecolor(self.STYLE_CONFIG['figure_facecolor'])
        wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=90,
                                          textprops={'color': self.STYLE_CONFIG['text_color'], 'fontsize': 11})
        for autotext in autotexts:
            autotext.set_color('white'); autotext.set_weight('bold'); autotext.set_fontsize(10)
        ax.set_title('Kategorilere Göre Yorum Dağılımı', fontsize=14, pad=20, color=self.STYLE_CONFIG['text_color'])
        ax.axis('equal')
        plt.tight_layout(pad=1.5)
        buffer = BytesIO()
        plt.savefig(buffer, format='png', bbox_inches='tight', dpi=100, facecolor=fig.get_facecolor())
        buffer.seek(0)
        plt.close(fig)
        return buffer

    def _create_category_timeseries_chart(self, category_name, category_csv_col):
        if category_csv_col not in self.df_comments.columns: return None
        df_cat = self.df_comments[(self.df_comments[category_csv_col] == 1) & (self.df_comments['hesaplanan_tarih'].notna())]
        if df_cat.empty: return None
        time_span_days = (df_cat['hesaplanan_tarih'].max() - df_cat['hesaplanan_tarih'].min()).days
        if time_span_days < 1: time_span_days = 1
        resample_period = 'M' if time_span_days > 90 else ('W' if time_span_days > 14 else 'D')
        date_format = '%Y-%b' if resample_period == 'M' else ('%Y-%b-%d' if resample_period == 'W' else '%b %d')
        df_indexed = df_cat.set_index('hesaplanan_tarih')
        resampled_pos = df_indexed[df_indexed['duygu_tahmini'] == 'positive']['duygu_skoru'].resample(resample_period).mean().dropna()
        resampled_neg = df_indexed[df_indexed['duygu_tahmini'] == 'negative']['duygu_skoru'].resample(resample_period).mean().dropna()
        if resampled_pos.empty and resampled_neg.empty: return None
        fig, ax = plt.subplots(figsize=(10, 5))
        self._apply_dark_theme(fig, ax)
        pos_linestyle = '-' if len(resampled_pos) > 1 else 'none'
        neg_linestyle = '-' if len(resampled_neg) > 1 else 'none'
        if not resampled_pos.empty:
            ax.plot(resampled_pos.index, resampled_pos.values, marker='o', linestyle=pos_linestyle, color=self.STYLE_CONFIG['positive_color'], label='Pozitif Ortalama')
        if not resampled_neg.empty:
            ax.plot(resampled_neg.index, resampled_neg.values, marker='o', linestyle=neg_linestyle, color=self.STYLE_CONFIG['negative_color'], label='Negatif Ortalama')
        ax.axhline(0, color=self.STYLE_CONFIG['grid_color'], linestyle='--', linewidth=0.8)
        period_tr = {'D': 'Günlük', 'W': 'Haftalık', 'M': 'Aylık'}.get(resample_period, '')
        ax.set_title(f"{category_name} - Duygu Skoru Değişimi ({period_tr})", fontsize=14)
        ax.set_ylabel('Ortalama Duygu Skoru')
        ax.legend(loc='best'); ax.set_ylim(-1.05, 1.05)
        ax.xaxis.set_major_formatter(mdates.DateFormatter(date_format))
        fig.autofmt_xdate()
        ax.grid(True, which='both', linestyle='--', linewidth=0.5, color=self.STYLE_CONFIG['grid_color'])
        plt.tight_layout()
        buffer = BytesIO()
        plt.savefig(buffer, format='png', dpi=150, facecolor=fig.get_facecolor())
        buffer.seek(0)
        plt.close(fig)
        return buffer

    def _ollama_gemma_chat(self, prompt_text):
        if ollama is None: return "Ollama kütüphanesi mevcut değil."
        try:
            response = ollama.generate(
                model=self.llm_model_name, prompt=prompt_text, options={'temperature': self.llm_temperature}
            )
            return response.get('response', '')
        except Exception as e:
            return f"Ollama hatası: {e}"

    def _parse_llm_section(self, response, section_key):
        start_tag = f"[{section_key.upper()}_BASLANGIC]"; end_tag = f"[{section_key.upper()}_BITIS]"
        start_idx = response.find(start_tag)
        if start_idx == -1: return f"LLM_PARSE_HATASI: '{section_key}' bölümü yanıtta bulunamadı."
        start_idx += len(start_tag)
        end_idx = response.find(end_tag, start_idx)
        if end_idx == -1: return f"LLM_PARSE_HATASI: '{section_key}' bölümünün bitiş etiketi bulunamadı."
        return response[start_idx:end_idx].strip()

    # --- YENİ EKLENEN YARDIMCI FONKSİYONLAR ---
    def _add_figure(self, doc, buffer, description, figure_counter):
        if buffer is None: return figure_counter
        try:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(buffer, width=Inches(6.0))
        except Exception as e:
            print(f"HATA: Resim eklenemedi - {e}")
            doc.add_paragraph(f"[HATA: Grafik yüklenemedi - {description}]", style='Emphasis')
            return figure_counter
        caption = f'Şekil {figure_counter}: {description}'
        p_caption = doc.add_paragraph(caption, style='Caption')
        p_caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p_caption.runs:
            run.font.size = Pt(9); run.font.italic = True; run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        return figure_counter + 1

    ### BU FONKSİYONU TAMAMEN DEĞİŞTİRİN ###

    def _setup_document_styles(self, doc):
        """Word belgesi için temel stilleri ve sayfa ayarlarını yapar."""
        # --- Sayfa Ayarları (Daha Sıkı Düzen) ---
        section = doc.sections[0]
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # --- Temel Normal Stil ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

        # --- Ana Başlık Stili (H1) ---
        style_h1 = doc.styles.add_style('RaporBaslik1', WD_STYLE_TYPE.PARAGRAPH)
        style_h1.base_style = doc.styles['Heading 1']
        font = style_h1.font
        font.name = 'Calibri'
        font.size = Pt(16)
        font.bold = True
        font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        p_fmt = style_h1.paragraph_format
        p_fmt.space_before = Pt(18) # Başlık öncesi boşluğu biraz artıralım
        p_fmt.space_after = Pt(6)
        
        # --- EN ÖNEMLİ DEĞİŞİKLİK: Otomatik sayfa sonunu kaldırıyoruz ---
        p_fmt.page_break_before = False  # Bu satırı True'dan False'a çevirdik!
        # --- DEĞİŞİKLİK SONU ---

        # --- Alt Başlık Stili (H2) ---
        style_h2 = doc.styles.add_style('RaporBaslik2', WD_STYLE_TYPE.PARAGRAPH)
        style_h2.base_style = doc.styles['Heading 2']
        font = style_h2.font
        font.name = 'Calibri'
        font.size = Pt(13)
        font.bold = True
        font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
        p_fmt = style_h2.paragraph_format
        p_fmt.space_before = Pt(12)
        p_fmt.space_after = Pt(4)

        # --- Grafik Yorum Stili ---
        style_comment = doc.styles.add_style('ChartComment', WD_STYLE_TYPE.PARAGRAPH)
        style_comment.base_style = doc.styles['Normal']
        font = style_comment.font
        font.name = 'Calibri'
        font.size = Pt(10)
        font.italic = True
        p_fmt = style_comment.paragraph_format
        p_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p_fmt.space_before = Pt(4)
        p_fmt.space_after = Pt(12) # Grafik sonrası boşluğu biraz artıralım

    def _get_llm_timeseries_comment(self, category_name, resampled_pos, resampled_neg):
        """Belirli bir kategori zaman serisi verisi için LLM'den zenginleştirilmiş yorum alır."""
        if ollama is None: return "Ollama kütüphanesi bulunamadı."
        
        prompt_data = f"Analiz Edilen Kategori: {category_name}\n"
        has_data = False
        
        if not resampled_pos.empty:
            has_data = True
            prompt_data += (f"- Pozitif Yorumlar: Başlangıç skoru ~{resampled_pos.values[0]:.2f}, "
                            f"bitiş skoru ~{resampled_pos.values[-1]:.2f}. "
                            f"Dönem içindeki en yüksek skor ~{resampled_pos.values.max():.2f}, "
                            f"en düşük skor ~{resampled_pos.values.min():.2f}.\n")
        if not resampled_neg.empty:
            has_data = True
            prompt_data += (f"- Negatif Yorumlar: Başlangıç skoru ~{resampled_neg.values[0]:.2f}, "
                            f"bitiş skoru ~{resampled_neg.values[-1]:.2f}. "
                            f"Dönem içindeki en yüksek skor ~{resampled_neg.values.max():.2f}, "
                            f"en düşük skor ~{resampled_neg.values.min():.2f}.\n")
        
        if not has_data:
            return f"{category_name} kategorisi için yeterli zaman serisi verisi bulunmamaktadır."

        ts_prompt = f"""
Bir ürün raporu için, aşağıdaki zaman serisi verilerini analiz ederek bu kategorideki duygu trendini yorumlayan, 1-2 cümleden oluşan akıcı bir analiz paragrafı yaz. Yorumun başına 'Trend Analizi' gibi bir başlık EKLEME. Sadece analizi içeren paragrafı yaz.

Örnek: 'Fiyat kategorisindeki olumlu yorumlar, başlangıçta yüksek bir seviyede seyrederken dönem ortasında belirgin bir düşüş yaşamış ve sonrasında daha durağan bir hale gelmiştir.'

VERİLER:
{prompt_data}
ANALİZ PARAGRAFI:
"""
        try:
            # Bu yorumlar kısa olduğu için etiket parse etmeye gerek yok, direkt cevabı alabiliriz.
            return self._ollama_gemma_chat(ts_prompt).strip()
        except Exception as e:
            print(f"[HATA] Zaman serisi yorumu alınamadı: {e}")
            return "Zaman serisi grafiği için yorum oluşturulurken bir hata oluştu."

    @pyqtSlot()
    def run(self):
        self._is_running = True
        self.status_message.emit("Rapor oluşturma işlemi başlatıldı...")
        if not self._load_and_preprocess_data():
            self._is_running = False
            return
        try:
            # Adım sayacı kur - sabit 10 adım
            # 1: sentiment pie, 2: category pie, 3: ts bundle,
            # 4: ozet, 5: duygu, 6: kategori, 7: sonuc, 8: oneriler,
            # 9: belge derleme, 10: kaydetme
            total_steps = 10 if not self.charts_only else 3
            current_step = 0

            # --- GRAFİK OLUŞTURMA VE ZAMAN SERİSİ YORUMLAMA ---
            # (Bu kısım öncekiyle aynı, sadece kopyalıyorum)
            self.status_message.emit("Grafikler oluşturuluyor...")
            sentiment_pie_buffer = self._create_sentiment_pie_chart()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)
            category_pie_buffer = self._create_category_pie_chart()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)
            category_ts_dict = {}; category_ts_comments_dict = {}
            all_categories = []
            available_category_cols = [col for col in self.df_comments.columns if col not in ['text', 'duygu_tahmini', 'duygu_skoru', 'hesaplanan_tarih', 'date'] and self.df_comments[col].dtype in ['int64', 'float64']]
            for cat in self.category_definitions:
                if cat['csv_col'] in self.df_comments.columns: all_categories.append(cat)
            for col in available_category_cols:
                if not any(cat['csv_col'] == col for cat in self.category_definitions):
                    all_categories.append({'display_name': col.replace('_', ' ').title(), 'csv_col': col})
            for cat in all_categories:
                if self._stop_requested: break
                ts_buffer = self._create_category_timeseries_chart(cat['display_name'], cat['csv_col'])
                if ts_buffer:
                    category_ts_dict[cat['display_name']] = ts_buffer
                    df_cat = self.df_comments[(self.df_comments[cat['csv_col']] == 1) & (self.df_comments['hesaplanan_tarih'].notna())]
                    if not df_cat.empty:
                        time_span_days = (df_cat['hesaplanan_tarih'].max() - df_cat['hesaplanan_tarih'].min()).days
                        resample_period = 'M' if time_span_days > 90 else ('W' if time_span_days > 14 else 'D')
                        df_indexed = df_cat.set_index('hesaplanan_tarih')
                        resampled_pos = df_indexed[df_indexed['duygu_tahmini'] == 'positive']['duygu_skoru'].resample(resample_period).mean().dropna()
                        resampled_neg = df_indexed[df_indexed['duygu_tahmini'] == 'negative']['duygu_skoru'].resample(resample_period).mean().dropna()
                        if not resampled_pos.empty or not resampled_neg.empty:
                            comment = self._get_llm_timeseries_comment(cat['display_name'], resampled_pos, resampled_neg)
                            category_ts_comments_dict[cat['display_name']] = comment
            
            # (Arayüze grafik gönderme kısmı da aynı)
            final_chart_buffers = {}
            if sentiment_pie_buffer: final_chart_buffers['sentiment_pie'] = sentiment_pie_buffer.getvalue()
            if category_pie_buffer: final_chart_buffers['category_pie'] = category_pie_buffer.getvalue()
            if category_ts_dict:
                for cat_name, buf in category_ts_dict.items():
                    final_chart_buffers[f"timeseries_{cat_name.replace(' ', '_').lower()}"] = buf.getvalue()
            if final_chart_buffers: self.charts_generated.emit(final_chart_buffers)
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 3: ts bundle

            if self._stop_requested or self.charts_only:
                self.finished.emit(""); return

            # --- LLM İÇİN TEMEL BİLGİ METNİ HAZIRLIĞI ---
            self.status_message.emit("LLM için analiz metinleri hazırlanıyor...")
            total_comments = len(self.df_comments); sentiment_counts = self.df_comments['duygu_tahmini'].value_counts()
            stats = {'total': total_comments, 'positive': sentiment_counts.get('positive', 0), 'negative': sentiment_counts.get('negative', 0), 'neutral': sentiment_counts.get('neutral', 0)}
            category_summary_list = []; example_comments_text = ""
            for cat in all_categories:
                count = int(self.df_comments[cat['csv_col']].sum())
                if count > 0:
                    category_summary_list.append(f"- {cat['display_name']}: {count} yorum")
                    df_cat = self.df_comments[self.df_comments[cat['csv_col']] == 1].copy()
                    df_cat_pos = df_cat[df_cat['duygu_tahmini'] == 'positive'].sort_values('duygu_skoru', ascending=False)
                    df_cat_neg = df_cat[df_cat['duygu_tahmini'] == 'negative'].sort_values('duygu_skoru', ascending=True)
                    example_comments_text += f"\n--- {cat['display_name']} Kategorisi Örnekleri ---\n"
                    if not df_cat_pos.empty: example_comments_text += f"En Pozitif Yorum: \"{df_cat_pos.iloc[0]['text'][:200]}...\"\n"
                    if not df_cat_neg.empty: example_comments_text += f"En Negatif Yorum: \"{df_cat_neg.iloc[0]['text'][:200]}...\"\n"
            category_summary = "\n".join(category_summary_list)
            
            base_context_prompt = f"""Bir ürün yorum analizi raporu için Türkçe metinler üreteceksin. Sana vereceğim genel verilere dayanarak, senden istediğim bölümü yaz.

--- GENEL VERİLER ---
Ürün Adı: {self.product_name}
Toplam Yorum Sayısı: {stats['total']}
Pozitif Yorum Sayısı: {stats['positive']} (%{((stats['positive']/stats['total']*100) if stats['total'] > 0 else 0):.1f})
Negatif Yorum Sayısı: {stats['negative']} (%{((stats['negative']/stats['total']*100) if stats['total'] > 0 else 0):.1f})
Nötr Yorum Sayısı: {stats['neutral']} (%{((stats['neutral']/stats['total']*100) if stats['total'] > 0 else 0):.1f})
Kategori Bazlı Yorum Sayıları:
{category_summary}
Örnek Yorumlar (Analizine derinlik katmak için kullan):
{example_comments_text}
"""
            # --- LLM'DEN BÖLÜMLERİ AYRI AYRI İSTE ---
            self.status_message.emit("LLM'den rapor metinleri alınıyor...")

            prompt_ozet = base_context_prompt + "\n--- İSTEK ---\nYukarıdaki genel verilere dayanarak, ürün hakkındaki kullanıcı yorumlarının genel bir özetini tek paragrafta yaz. Sadece istenen paragrafı yaz, başka bir şey ekleme."
            ozet_text = self._ollama_gemma_chat(prompt_ozet).strip()
            if not self._is_running: return
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 4

            prompt_duygu = base_context_prompt + "\n--- İSTEK ---\nYukarıdaki duygu analizi sayılarına bakarak, duygu dağılımı grafiğini yorumlayan tek paragraflık bir metin yaz. Sadece istenen paragrafı yaz."
            duygu_grafik_yorum_text = self._ollama_gemma_chat(prompt_duygu).strip()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 5

            prompt_kategori = base_context_prompt + "\n--- İSTEK ---\nYukarıdaki kategori bazlı yorum sayılarına ve örnek yorumlara dayanarak, kategori dağılımını ve öne çıkan temaları yorumlayan tek paragraflık bir metin yaz. Sadece istenen paragrafı yaz."
            kategori_grafik_yorum_text = self._ollama_gemma_chat(prompt_kategori).strip()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 6

            prompt_sonuc = base_context_prompt + "\n--- İSTEK ---\nTüm bu verilere dayanarak, ürün için genel bir sonuç ve değerlendirme paragrafı yaz. Sadece istenen paragrafı yaz."
            sonuc_degerlendirme_text = self._ollama_gemma_chat(prompt_sonuc).strip()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 7

            prompt_oneriler = base_context_prompt + "\n--- İSTEK ---\nÖzellikle negatif yorumlardaki ana temaları dikkate alarak, ürün veya pazarlama stratejisini iyileştirmek için 2-3 maddelik somut ve eyleme geçirilebilir öneriler sun. Sadece istenen maddeleri yaz."
            oneriler_text = self._ollama_gemma_chat(prompt_oneriler).strip()
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 8

            # Artık _parse_llm_section fonksiyonuna ihtiyacımız kalmadı, çünkü her cevap zaten istediğimiz metin.

            # --- WORD RAPORU OLUŞTURMA (Önceki düzeltmeler dahil) ---
            self.status_message.emit("Word raporu oluşturuluyor...")
            doc = Document()
            self._setup_document_styles(doc) # Yeni stil ayarlarımız burada uygulanacak.
            figure_counter = 1
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 9

            # --- ANA BAŞLIK ---
            doc.add_heading(f'{self.product_name} - Yorum Analizi Raporu', level=0)
            doc.add_paragraph(f"Rapor Tarihi: {datetime.now().strftime('%d.%m.%Y')}")
            doc.add_paragraph() # Kısa boşluk

            # --- GENEL ÖZET ---
            h1 = doc.add_heading('Genel Özet', level=1)
            h1.style = 'RaporBaslik1'
            doc.add_paragraph(ozet_text or "Bu bölüm için metin üretilemedi.")

            # --- DUYGU ANALİZİ ---
            h1 = doc.add_heading('Duygu Analizi Sonuçları', level=1)
            h1.style = 'RaporBaslik1'
            doc.add_paragraph("Aşağıdaki tablo, yorumların duygu dağılımını sayısal ve oransal olarak özetlemektedir.")
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Duygu Durumu'; hdr_cells[1].text = 'Yorum Sayısı'; hdr_cells[2].text = 'Oran (%)'
            for cell in hdr_cells: cell.paragraphs[0].runs[0].font.bold = True
            stats = {'total': len(self.df_comments), 'positive': self.df_comments['duygu_tahmini'].value_counts().get('positive', 0), 'negative': self.df_comments['duygu_tahmini'].value_counts().get('negative', 0), 'neutral': self.df_comments['duygu_tahmini'].value_counts().get('neutral', 0)}
            for duygu, sayi in [('Pozitif', stats.get('positive', 0)), ('Negatif', stats.get('negative', 0)), ('Nötr', stats.get('neutral', 0))]:
                row_cells = table.add_row().cells
                row_cells[0].text = duygu
                row_cells[1].text = str(sayi)
                oran = (sayi / stats['total'] * 100) if stats['total'] > 0 else 0
                row_cells[2].text = f'{oran:.1f}%'
            doc.add_paragraph()
            figure_counter = self._add_figure(doc, sentiment_pie_buffer, 'Yorumların Genel Duygu Dağılımı', figure_counter)
            doc.add_paragraph(duygu_grafik_yorum_text or "Bu bölüm için metin üretilemedi.")
            # doc.add_page_break() # <-- DEĞİŞİKLİK: KALDIRILDI.

            # --- KATEGORİ ANALİZİ ---
            h1 = doc.add_heading('Kategori Analizi Sonuçları', level=1)
            h1.style = 'RaporBaslik1'
            doc.add_paragraph("Aşağıdaki tablo, yorumların hangi kategorilerde yoğunlaştığını göstermektedir.")

            # Kategori listesini oluştur
            all_categories = []
            available_category_cols = [col for col in self.df_comments.columns
                                     if col not in ['text', 'duygu_tahmini', 'duygu_skoru', 'hesaplanan_tarih', 'date']
                                     and self.df_comments[col].dtype in ['int64', 'float64']]
            for cat in self.category_definitions:
                if cat['csv_col'] in self.df_comments.columns:
                    all_categories.append(cat)
            for col in available_category_cols:
                if not any(cat['csv_col'] == col for cat in self.category_definitions):
                    all_categories.append({'display_name': col.replace('_', ' ').title(), 'csv_col': col})

            # Kategori tablosunu oluştur
            cat_table = doc.add_table(rows=1, cols=2)
            cat_table.style = 'Table Grid'
            cat_hdr_cells = cat_table.rows[0].cells
            cat_hdr_cells[0].text = 'Kategori'
            cat_hdr_cells[1].text = 'Yorum Sayısı'
            for cell in cat_hdr_cells:
                cell.paragraphs[0].runs[0].font.bold = True

            # Sıralanmış kategorileri ekle
            sorted_categories = sorted(all_categories, key=lambda x: self.df_comments[x['csv_col']].sum(), reverse=True)
            for cat in sorted_categories:
                count = int(self.df_comments[cat['csv_col']].sum())
                if count > 0:
                    row_cells = cat_table.add_row().cells
                    row_cells[0].text = cat['display_name']
                    row_cells[1].text = str(count)

            doc.add_paragraph() # Kısa boşluk
            figure_counter = self._add_figure(doc, category_pie_buffer, 'Yorumların Kategorilere Göre Dağılımı', figure_counter)
            doc.add_paragraph(kategori_grafik_yorum_text or "Bu bölüm için metin üretilemedi.")

            # --- KATEGORİ ZAMAN SERİSİ ANALİZİ ---
            if category_ts_dict:
                h2 = doc.add_heading('Kategori Bazında Duygu Değişimi', level=2)
                h2.style = 'RaporBaslik2'
                doc.add_paragraph("Aşağıdaki grafikler, önemli kategorilerdeki olumlu ve olumsuz yorumların duygu skorlarının zaman içindeki değişimini göstermektedir.")

                for cat_name, ts_buffer in category_ts_dict.items():
                    figure_counter = self._add_figure(doc, ts_buffer, f'{cat_name} Kategorisindeki Duygu Skorlarının Zamana Göre Değişimi', figure_counter)

                    # Yorum ekleme
                    if cat_name in category_ts_comments_dict:
                        comment_text = category_ts_comments_dict[cat_name]
                        doc.add_paragraph(comment_text, style='ChartComment')
                    else:
                        doc.add_paragraph("Bu grafik için otomatik yorum oluşturulamadı.", style='ChartComment')

            # --- SONUÇ VE ÖNERİLER ---
            h1 = doc.add_heading('Sonuç ve Değerlendirme', level=1)
            h1.style = 'RaporBaslik1'
            doc.add_paragraph(sonuc_degerlendirme_text or "Bu bölüm için metin üretilemedi.")

            h1 = doc.add_heading('Öneriler', level=1)
            h1.style = 'RaporBaslik1'
            doc.add_paragraph(oneriler_text or "Bu bölüm için metin üretilemedi.")
            
            # --- KAYDETME ---
            output_path = os.path.join(OUTPUT_DIR, REPORT_FILE_TEMPLATE.format(product_id=self.product_id))
            os.makedirs(OUTPUT_DIR, exist_ok=True)
            doc.save(output_path)
            current_step += 1; self.progress_updated.emit(current_step, total_steps)  # 10
            self.finished.emit(output_path)

        except Exception as e:
            import traceback
            error_msg = f"Rapor oluşturma hatası: {e}"
            print(f"[HATA] ReportBuilder run() hatası: {str(e)}")
            print(f"[HATA] Stack trace:\n{traceback.format_exc()}")
            self.error.emit(error_msg)
        finally:
            self._is_running = False